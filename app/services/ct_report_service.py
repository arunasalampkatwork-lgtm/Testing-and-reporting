from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches

from PySide6.QtWidgets import (
    QFileDialog,
    QMessageBox,
)


class CTReportService:

    def __init__(
        self,
        project_folder
    ):

        self.project_folder = Path(
            project_folder
        )

    # =====================================================
    # VALUE CHECK
    # =====================================================

    @staticmethod
    def has_value(value):

        if value is None:
            return False

        if isinstance(value, str):

            return bool(
                value.strip()
            )

        return True

    # =====================================================
    # GENERATE
    # =====================================================

    def generate_report(
        self,
        record,
        parent=None
    ):

        default_name = (
            f"CT_Test_"
            f"{record.get('test_id', 'Report')}"
            f".docx"
        )

        output_path, _ = (
            QFileDialog.getSaveFileName(

                parent,

                "Save CT Test Report",

                str(
                    self.project_folder
                    /
                    "reports"
                    /
                    default_name
                ),

                "Word Document (*.docx)"
            )
        )

        if not output_path:
            return None

        output_path = Path(
            output_path
        )

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True
        )

        document = Document()

        self.configure_document(
            document
        )

        measurements = (
            record.get(
                "measurements",
                {}
            )
            or {}
        )

        # =================================================
        # TITLE
        # =================================================

        self.add_title(
            document,
            "CURRENT TRANSFORMER TEST REPORT"
        )

        # =================================================
        # TEST INFORMATION
        # =================================================

        test_information = [

            (
                "Test ID",
                record.get(
                    "test_id",
                    ""
                )
            ),

            (
                "Test Date",
                record.get(
                    "test_date",
                    ""
                )
            ),

            (
                "Project ID",
                record.get(
                    "project_id",
                    ""
                )
            ),

            (
                "Panel ID",
                record.get(
                    "panel_id",
                    ""
                )
            ),

            (
                "Component ID",
                record.get(
                    "component_id",
                    ""
                )
            ),

            (
                "Test Type",
                record.get(
                    "test_type",
                    "CT"
                )
            ),

        ]

        if measurements.get(
            "is_three_phase",
            False
        ):

            test_information.append(
                (
                    "3-Phase CT",
                    "Yes"
                )
            )

        self.add_optional_info_section(
            document,
            "TEST INFORMATION",
            test_information
        )

        # =================================================
        # CT DETAILS
        # =================================================

        ct_details = [

            (
                "CT",
                measurements.get(
                    "ct_name",
                    ""
                )
            ),

            (
                "Manufacturer",
                measurements.get(
                    "manufacturer",
                    ""
                )
            ),

            (
                "Model",
                measurements.get(
                    "model",
                    ""
                )
            ),

            (
                "Serial Number",
                measurements.get(
                    "serial_number",
                    ""
                )
            ),

            (
                "CT Primary",
                measurements.get(
                    "ct_primary",
                    ""
                )
            ),

            (
                "CT Secondary",
                measurements.get(
                    "ct_secondary",
                    ""
                )
            ),

            (
                "CT Ratio",
                measurements.get(
                    "ct_ratio",
                    ""
                )
            ),

            (
                "Core",
                measurements.get(
                    "core",
                    ""
                )
            ),

            (
                "Class",
                measurements.get(
                    "ct_class",
                    ""
                )
            ),

            (
                "Burden",
                measurements.get(
                    "burden",
                    ""
                )
            ),

        ]

        self.add_optional_info_section(
            document,
            "CT DETAILS",
            ct_details
        )

        # =================================================
        # RATIO TEST
        # =================================================

        phase_tests = (
            measurements.get(
                "phase_tests",
                []
            )
            or []
        )

        # -------------------------------------------------
        # BACKWARD COMPATIBILITY
        # -------------------------------------------------

        if not phase_tests:

            old_ratio_values = [

                measurements.get(
                    "primary_current",
                    ""
                ),

                measurements.get(
                    "secondary_current",
                    ""
                ),

                measurements.get(
                    "measured_ratio",
                    ""
                ),

                measurements.get(
                    "ratio_error",
                    ""
                ),

            ]

            # Only create the legacy row if at least
            # one actual value exists.

            if any(
                self.has_value(value)
                for value in old_ratio_values
            ):

                phase_tests = [

                    {
                        "phase": "R",

                        "primary_current":
                            measurements.get(
                                "primary_current",
                                ""
                            ),

                        "secondary_current":
                            measurements.get(
                                "secondary_current",
                                ""
                            ),

                        "measured_ratio":
                            measurements.get(
                                "measured_ratio",
                                ""
                            ),

                        "ratio_error":
                            measurements.get(
                                "ratio_error",
                                ""
                            ),
                    }
                ]

        # -------------------------------------------------
        # ONLY SHOW RATIO SECTION IF DATA EXISTS
        # -------------------------------------------------

        valid_phase_tests = []

        for phase_data in phase_tests:

            if not isinstance(
                phase_data,
                dict
            ):
                continue

            values = [

                phase_data.get(
                    "phase",
                    ""
                ),

                phase_data.get(
                    "primary_current",
                    ""
                ),

                phase_data.get(
                    "secondary_current",
                    ""
                ),

                phase_data.get(
                    "measured_ratio",
                    ""
                ),

                phase_data.get(
                    "ratio_error",
                    ""
                ),

            ]

            if any(
                self.has_value(value)
                for value in values
            ):

                valid_phase_tests.append(
                    phase_data
                )

        if valid_phase_tests:

            self.add_heading(
                document,
                "RATIO TEST"
            )

            table = document.add_table(
                rows=1,
                cols=5
            )

            table.style = (
                "Table Grid"
            )

            table.alignment = (
                WD_TABLE_ALIGNMENT.CENTER
            )

            headers = [

                "Phase",

                "Injected Primary (A)",

                "Recorded Secondary (A)",

                "Measured Ratio",

                "Ratio Error (%)",

            ]

            for index, text in enumerate(
                headers
            ):

                cell = (
                    table.rows[0]
                    .cells[index]
                )

                cell.text = text

                for run in (
                    cell
                    .paragraphs[0]
                    .runs
                ):

                    run.bold = True

            for phase_data in valid_phase_tests:

                cells = (
                    table
                    .add_row()
                    .cells
                )

                values = [

                    phase_data.get(
                        "phase",
                        ""
                    ),

                    phase_data.get(
                        "primary_current",
                        ""
                    ),

                    phase_data.get(
                        "secondary_current",
                        ""
                    ),

                    phase_data.get(
                        "measured_ratio",
                        ""
                    ),

                    phase_data.get(
                        "ratio_error",
                        ""
                    ),

                ]

                for index, value in enumerate(
                    values
                ):

                    cells[index].text = (
                        ""
                        if value is None
                        else str(value)
                    )

        # =================================================
        # POLARITY
        # =================================================

        polarity = [

            (
                "Expected Polarity",
                measurements.get(
                    "expected_polarity",
                    ""
                )
            ),

            (
                "Observed Polarity",
                measurements.get(
                    "observed_polarity",
                    ""
                )
            ),

            (
                "Polarity Result",
                measurements.get(
                    "polarity_result",
                    ""
                )
            ),

        ]

        self.add_optional_info_section(
            document,
            "POLARITY",
            polarity
        )

        # =================================================
        # WINDING RESISTANCE
        # =================================================

        winding = [

            (
                "Phase R",
                measurements.get(
                    "resistance_phase_a",
                    ""
                )
            ),

            (
                "Phase Y",
                measurements.get(
                    "resistance_phase_b",
                    ""
                )
            ),

            (
                "Phase B",
                measurements.get(
                    "resistance_phase_c",
                    ""
                )
            ),

        ]

        self.add_optional_info_section(
            document,
            "CT WINDING RESISTANCE",
            winding
        )

        # =================================================
        # INSULATION
        # =================================================

        insulation = [

            (
                "Primary - Earth",
                measurements.get(
                    "ir_primary_earth",
                    ""
                )
            ),

            (
                "Secondary - Earth",
                measurements.get(
                    "ir_secondary_earth",
                    ""
                )
            ),

            (
                "Primary - Secondary",
                measurements.get(
                    "ir_primary_secondary",
                    ""
                )
            ),

            (
                "Test Voltage",
                measurements.get(
                    "ir_test_voltage",
                    ""
                )
            ),

            (
                "Test Duration",
                measurements.get(
                    "ir_test_duration",
                    ""
                )
            ),

        ]

        self.add_optional_info_section(
            document,
            "INSULATION RESISTANCE",
            insulation
        )

        # =================================================
        # KNEE POINT
        # =================================================

        knee_point = [

            (
                "Knee Point Voltage",
                measurements.get(
                    "knee_point_voltage",
                    ""
                )
            ),

            (
                "Knee Point Current",
                measurements.get(
                    "knee_point_current",
                    ""
                )
            ),

            (
                "Excitation Test Voltage",
                measurements.get(
                    "excitation_test_voltage",
                    ""
                )
            ),

            (
                "Excitation Test Current",
                measurements.get(
                    "excitation_test_current",
                    ""
                )
            ),

        ]

        self.add_optional_info_section(
            document,
            "KNEE POINT / EXCITATION",
            knee_point
        )

        # =================================================
        # BURDEN
        # =================================================

        burden = [

            (
                "Burden Test Current",
                measurements.get(
                    "burden_test_current",
                    ""
                )
            ),

            (
                "Measured Burden",
                measurements.get(
                    "measured_burden",
                    ""
                )
            ),

            (
                "Burden Error",
                measurements.get(
                    "burden_error",
                    ""
                )
            ),

        ]

        self.add_optional_info_section(
            document,
            "BURDEN",
            burden
        )

        # =================================================
        # RESULT
        # =================================================

        result = record.get(
            "result",
            ""
        )

        if self.has_value(result):

            self.add_heading(
                document,
                "TEST RESULT"
            )

            result_paragraph = (
                document.add_paragraph()
            )

            run = (
                result_paragraph
                .add_run(
                    str(result)
                )
            )

            run.bold = True

            run.font.size = Pt(
                13
            )

        # =================================================
        # REMARKS
        # =================================================

        remarks = (

            record.get(
                "remarks",
                ""
            )

            or

            measurements.get(
                "remarks",
                ""
            )

            or ""

        )

        if self.has_value(
            remarks
        ):

            self.add_heading(
                document,
                "REMARKS"
            )

            document.add_paragraph(
                str(remarks)
            )

        # =================================================
        # SIGNATURE
        # =================================================

        self.add_heading(
            document,
            "TESTING / APPROVAL"
        )

        table = document.add_table(
            rows=3,
            cols=2
        )

        table.style = (
            "Table Grid"
        )

        table.rows[0].cells[0].text = (
            "Tested By"
        )

        table.rows[1].cells[0].text = (
            "Reviewed By"
        )

        table.rows[2].cells[0].text = (
            "Date / Signature"
        )

        document.save(
            str(
                output_path
            )
        )

        QMessageBox.information(
            parent,
            "Report Generated",
            (
                "CT test report generated "
                "successfully.\n\n"
                f"{output_path}"
            )
        )

        return output_path

    # =====================================================
    # OPTIONAL INFO SECTION
    # =====================================================

    def add_optional_info_section(
        self,
        document,
        heading,
        rows
    ):

        valid_rows = []

        for label, value in rows:

            if self.has_value(
                value
            ):

                valid_rows.append(
                    (
                        label,
                        value
                    )
                )

        # Nothing was entered.
        # Don't even create the heading.

        if not valid_rows:
            return

        self.add_heading(
            document,
            heading
        )

        self.add_info_table(
            document,
            valid_rows
        )

    # =====================================================
    # DOCUMENT HELPERS
    # =====================================================

    @staticmethod
    def configure_document(
        document
    ):

        section = (
            document.sections[0]
        )

        section.top_margin = (
            Inches(0.6)
        )

        section.bottom_margin = (
            Inches(0.6)
        )

        section.left_margin = (
            Inches(0.65)
        )

        section.right_margin = (
            Inches(0.65)
        )

        document.styles[
            "Normal"
        ].font.name = "Arial"

        document.styles[
            "Normal"
        ].font.size = Pt(9)

    @staticmethod
    def add_title(
        document,
        title
    ):

        paragraph = (
            document.add_paragraph()
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = (
            paragraph
            .add_run(
                title
            )
        )

        run.bold = True

        run.font.size = Pt(
            17
        )

    @staticmethod
    def add_heading(
        document,
        text
    ):

        paragraph = (
            document.add_paragraph()
        )

        run = (
            paragraph
            .add_run(
                text
            )
        )

        run.bold = True

        run.font.size = Pt(
            12
        )

    @staticmethod
    def add_info_table(
        document,
        rows
    ):

        # ---------------------------------------------
        # FINAL SAFETY FILTER
        # ---------------------------------------------

        rows = [

            (
                label,
                value
            )

            for label, value in rows

            if (

                value is not None

                and

                (
                    not isinstance(
                        value,
                        str
                    )

                    or

                    bool(
                        value.strip()
                    )
                )
            )

        ]

        if not rows:
            return

        table = document.add_table(
            rows=1,
            cols=2
        )

        table.style = (
            "Table Grid"
        )

        table.rows[0].cells[0].text = (
            "Parameter"
        )

        table.rows[0].cells[1].text = (
            "Value"
        )

        for cell in (
            table.rows[0]
            .cells
        ):

            for run in (
                cell
                .paragraphs[0]
                .runs
            ):

                run.bold = True

        for label, value in rows:

            cells = (
                table
                .add_row()
                .cells
            )

            cells[0].text = (
                str(label)
            )

            cells[1].text = (
                str(value)
            )