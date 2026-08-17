from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PySide6.QtWidgets import QFileDialog, QMessageBox


class PanelReportService:
    """
    Generates a Word report for one panel.

    The service is intentionally self-contained. Helper methods that do not
    require instance state are static methods, so there is no accidental
    use of undefined `self` or `cls`.
    """

    def __init__(self, project_folder=None):
        self.project_folder = (
            Path(project_folder)
            if project_folder
            else Path.cwd()
        )

    # =====================================================
    # PUBLIC ENTRY POINT
    # =====================================================

    def generate_report(
        self,
        panel,
        components=None,
        protection_tests=None,
        component_tests=None,
        report_date=None,
        substation_name="",
        switchboard_name="",
        parent=None,
        output_path=None,
    ):
        """
        Generate a panel test report.

        `report_date` is expected to represent the selected test date.
        Filtering should normally already be performed by AssetView, but
        this method also filters supplied tests when a date is provided.
        """

        components = list(
            components or []
        )

        protection_tests = list(
            protection_tests or []
        )

        component_tests = list(
            component_tests or []
        )

        selected_date = self.normalize_date(
            report_date
        )

        if selected_date:
            protection_tests = [
                test
                for test in protection_tests
                if self.date_only(
                    self.get_value(
                        test,
                        "test_date",
                        ""
                    )
                ) == selected_date
            ]

            component_tests = [
                test
                for test in component_tests
                if self.date_only(
                    self.get_value(
                        test,
                        "test_date",
                        ""
                    )
                ) == selected_date
            ]

        panel_name = self.get_value(
            panel,
            "name",
            "Panel"
        )

        if not substation_name:
            substation_name = self.find_parent_name(
                panel,
                (
                    "substation",
                    "substation_name",
                )
            )

        if not switchboard_name:
            switchboard_name = self.find_parent_name(
                panel,
                (
                    "switchboard",
                    "switchboard_name",
                    "swbd",
                    "switchboard_no",
                )
            )

        report_date_text = (
            selected_date
            or
            self.normalize_date(
                datetime.now()
            )
        )

        # -------------------------------------------------
        # FILE NAME
        # -------------------------------------------------

        filename_parts = [
            self.safe_filename_part(
                substation_name
            ),
            self.safe_filename_part(
                switchboard_name
            ),
            self.safe_filename_part(
                panel_name
            ),
            self.safe_filename_part(
                report_date_text
            ),
        ]

        filename_parts = [
            part
            for part in filename_parts
            if part
        ]

        if not filename_parts:
            filename_parts = [
                "Panel",
                "Test Report",
                report_date_text,
            ]

        suggested_filename = (
            " - ".join(filename_parts)
            + ".docx"
        )

        if output_path is None:
            reports_folder = (
                self.project_folder
                / "reports"
            )

            reports_folder.mkdir(
                parents=True,
                exist_ok=True
            )

            default_path = (
                reports_folder
                / suggested_filename
            )

            output_path, _ = (
                QFileDialog.getSaveFileName(
                    parent,
                    "Save Panel Test Report",
                    str(default_path),
                    "Word Document (*.docx)",
                )
            )

            if not output_path:
                return None

        output_path = Path(
            output_path
        )

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True
        )

        # -------------------------------------------------
        # DOCUMENT
        # -------------------------------------------------

        document = Document()

        self.configure_document(
            document
        )

        self.add_title(
            document,
            "PANEL TEST REPORT"
        )

        # -------------------------------------------------
        # PANEL DETAILS
        # -------------------------------------------------

        self.add_heading(
            document,
            "PANEL DETAILS"
        )

        panel_rows = [
            (
                "Substation",
                substation_name
            ),
            (
                "Switchboard",
                switchboard_name
            ),
            (
                "Panel No.",
                panel_name
            ),
            (
                "Test Date",
                report_date_text
            ),
            (
                "Feed Equipment",
                self.get_value(
                    panel,
                    "equipment_name",
                    ""
                )
            ),
            (
                "Equipment Type",
                self.get_value(
                    panel,
                    "equipment_type",
                    ""
                )
            ),
            (
                "Number of CTs",
                self.get_value(
                    panel,
                    "ct_count",
                    ""
                )
            ),
            (
                "Numerical Relays",
                self.get_value(
                    panel,
                    "relay_count",
                    ""
                )
            ),
            (
                "Auxiliary Relays",
                self.get_value(
                    panel,
                    "aux_count",
                    ""
                )
            ),
            (
                "Meters",
                self.get_value(
                    panel,
                    "meter_count",
                    ""
                )
            ),
        ]

        self.add_key_value_table(
            document,
            self.remove_empty_rows(
                panel_rows
            )
        )

        # -------------------------------------------------
        # COMPONENTS
        # -------------------------------------------------

        if components:

            self.add_heading(
                document,
                "INSTALLED TEST COMPONENTS"
            )

            self.add_components_table(
                document,
                components
            )

        # -------------------------------------------------
        # COMPONENT NAME LOOKUP
        # -------------------------------------------------

        component_names = {}

        for component in components:

            component_id = self.get_value(
                component,
                "component_id",
                ""
            )

            component_name = self.get_value(
                component,
                "name",
                ""
            )

            if component_id:
                component_names[
                    str(component_id)
                ] = (
                    component_name
                    or
                    component_id
                )

        # -------------------------------------------------
        # SUMMARY
        # -------------------------------------------------

        if (
            protection_tests
            or
            component_tests
        ):

            self.add_heading(
                document,
                "TEST SUMMARY"
            )

            self.add_summary_table(
                document,
                protection_tests,
                component_tests,
                component_names
            )

        # -------------------------------------------------
        # COMPONENT DETAILS
        # -------------------------------------------------

        if component_tests:

            self.add_heading(
                document,
                "COMPONENT TEST DETAILS"
            )

            for test in component_tests:

                self.add_component_test_detail(
                    document,
                    test,
                    component_names
                )

        # -------------------------------------------------
        # PROTECTION DETAILS
        # -------------------------------------------------

        if protection_tests:

            self.add_heading(
                document,
                "PROTECTION TEST DETAILS"
            )

            for test in protection_tests:

                self.add_protection_test_detail(
                    document,
                    test,
                    component_names
                )

        # -------------------------------------------------
        # OVERALL RESULT
        # -------------------------------------------------

        self.add_heading(
            document,
            "OVERALL PANEL RESULT"
        )

        overall = self.calculate_overall_result(
            protection_tests,
            component_tests
        )

        self.add_result(
            document,
            overall
        )

        # -------------------------------------------------
        # SIGNATURES
        # -------------------------------------------------

        self.add_heading(
            document,
            "TESTING / APPROVAL"
        )

        signature_table = document.add_table(
            rows=3,
            cols=2
        )

        signature_table.style = "Table Grid"

        signature_table.cell(
            0, 0
        ).text = "Tested By"

        signature_table.cell(
            1, 0
        ).text = "Reviewed By"

        signature_table.cell(
            2, 0
        ).text = "Date / Signature"

        document.save(
            str(output_path)
        )

        if parent is not None:
            QMessageBox.information(
                parent,
                "Report Generated",
                (
                    "Panel test report generated "
                    "successfully.\n\n"
                    f"{output_path}"
                )
            )

        return output_path

    # =====================================================
    # COMPONENT TABLE
    # =====================================================

    @staticmethod
    def add_components_table(
        document,
        components
    ):

        table = document.add_table(
            rows=1,
            cols=5
        )

        table.style = "Table Grid"

        headers = [
            "Component",
            "Type",
            "Manufacturer",
            "Model",
            "Serial Number",
        ]

        for index, header in enumerate(
            headers
        ):
            table.cell(
                0,
                index
            ).text = header

        for component in components:

            values = [
                PanelReportService.get_value(
                    component,
                    "name",
                    ""
                ),
                PanelReportService.get_value(
                    component,
                    "component_type",
                    ""
                ),
                PanelReportService.get_value(
                    component,
                    "manufacturer",
                    ""
                ),
                PanelReportService.get_value(
                    component,
                    "model",
                    ""
                ),
                PanelReportService.get_value(
                    component,
                    "serial_number",
                    ""
                ),
            ]

            if all(
                PanelReportService.is_empty_value(
                    value
                )
                for value in values
            ):
                continue

            cells = table.add_row().cells

            for index, value in enumerate(
                values
            ):
                cells[index].text = (
                    ""
                    if PanelReportService.is_empty_value(
                        value
                    )
                    else str(value)
                )

    # =====================================================
    # SUMMARY
    # =====================================================

    @staticmethod
    def add_summary_table(
        document,
        protection_tests,
        component_tests,
        component_names
    ):

        table = document.add_table(
            rows=1,
            cols=5
        )

        table.style = "Table Grid"

        headers = [
            "Test ID",
            "Date",
            "Test Type",
            "Component",
            "Result",
        ]

        for index, header in enumerate(
            headers
        ):
            table.cell(
                0,
                index
            ).text = header

        for test in component_tests:

            component_id = (
                PanelReportService.get_value(
                    test,
                    "component_id",
                    ""
                )
            )

            values = [
                PanelReportService.get_value(
                    test,
                    "test_id",
                    ""
                ),
                PanelReportService.date_only(
                    PanelReportService.get_value(
                        test,
                        "test_date",
                        ""
                    )
                ),
                PanelReportService.get_value(
                    test,
                    "test_type",
                    ""
                ),
                component_names.get(
                    str(component_id),
                    component_id
                ),
                PanelReportService.get_value(
                    test,
                    "result",
                    ""
                ),
            ]

            PanelReportService.add_summary_row(
                table,
                values
            )

        for test in protection_tests:

            relay_id = (
                PanelReportService.get_value(
                    test,
                    "relay_id",
                    ""
                )
            )

            values = [
                PanelReportService.get_value(
                    test,
                    "test_id",
                    ""
                ),
                PanelReportService.date_only(
                    PanelReportService.get_value(
                        test,
                        "test_date",
                        ""
                    )
                ),
                PanelReportService.get_value(
                    test,
                    "protection_code",
                    ""
                ),
                component_names.get(
                    str(relay_id),
                    relay_id
                ),
                PanelReportService.get_value(
                    test,
                    "result",
                    ""
                ),
            ]

            PanelReportService.add_summary_row(
                table,
                values
            )

    # =====================================================
    # COMPONENT DETAIL
    # =====================================================

    @staticmethod
    def add_component_test_detail(
        document,
        test,
        component_names
    ):

        component_id = (
            PanelReportService.get_value(
                test,
                "component_id",
                ""
            )
        )

        component_name = component_names.get(
            str(component_id),
            component_id
        )

        test_type = (
            PanelReportService.get_value(
                test,
                "test_type",
                ""
            )
        )

        heading = str(
            component_name
            or
            component_id
            or
            "Component"
        )

        if test_type:
            heading += (
                f" - {test_type}"
            )

        PanelReportService.add_heading(
            document,
            heading
        )

        measurements = (
            PanelReportService.get_value(
                test,
                "measurements",
                {}
            )
            or
            {}
        )

        if not isinstance(
            measurements,
            dict
        ):
            measurements = {}

        scalar_rows = []

        for key, value in measurements.items():

            if key in (
                "phase_tests",
                "functions",
            ):
                continue

            if isinstance(
                value,
                (dict, list, tuple)
            ):
                continue

            if PanelReportService.is_empty_value(
                value
            ):
                continue

            scalar_rows.append(
                (
                    PanelReportService.pretty_label(
                        key
                    ),
                    value
                )
            )

        PanelReportService.add_key_value_table(
            document,
            scalar_rows
        )

        phase_tests = measurements.get(
            "phase_tests"
        )

        if phase_tests:

            PanelReportService.add_heading(
                document,
                "PHASE-WISE CT TEST RESULTS"
            )

            PanelReportService.add_phase_test_table(
                document,
                phase_tests
            )

        functions = measurements.get(
            "functions"
        )

        if functions:

            PanelReportService.add_heading(
                document,
                "METER MEASUREMENT RESULTS"
            )

            PanelReportService.add_meter_function_table(
                document,
                functions
            )

        result = PanelReportService.get_value(
            test,
            "result",
            ""
        )

        if not PanelReportService.is_empty_value(
            result
        ):

            PanelReportService.add_result(
                document,
                result
            )

        remarks = PanelReportService.get_value(
            test,
            "remarks",
            ""
        )

        if not PanelReportService.is_empty_value(
            remarks
        ):

            document.add_paragraph(
                f"Remarks: {remarks}"
            )

    # =====================================================
    # PROTECTION DETAIL
    # =====================================================

    @staticmethod
    def add_protection_test_detail(
        document,
        test,
        component_names
    ):

        relay_id = (
            PanelReportService.get_value(
                test,
                "relay_id",
                ""
            )
        )

        relay_name = component_names.get(
            str(relay_id),
            relay_id
        )

        protection_code = (
            PanelReportService.get_value(
                test,
                "protection_code",
                ""
            )
        )

        heading = str(
            relay_name
            or
            relay_id
            or
            "Relay"
        )

        if protection_code:
            heading += (
                f" - {protection_code}"
            )

        PanelReportService.add_heading(
            document,
            heading
        )

        measurements = (
            PanelReportService.get_value(
                test,
                "measurements",
                {}
            )
            or
            {}
        )

        if not isinstance(
            measurements,
            dict
        ):
            measurements = {}

        rows = []

        for key, value in measurements.items():

            if isinstance(
                value,
                (dict, list, tuple)
            ):
                continue

            if PanelReportService.is_empty_value(
                value
            ):
                continue

            rows.append(
                (
                    PanelReportService.pretty_label(
                        key
                    ),
                    value
                )
            )

        PanelReportService.add_key_value_table(
            document,
            rows
        )

        result = PanelReportService.get_value(
            test,
            "result",
            ""
        )

        if not PanelReportService.is_empty_value(
            result
        ):

            PanelReportService.add_result(
                document,
                result
            )

        remarks = PanelReportService.get_value(
            test,
            "remarks",
            ""
        )

        if not PanelReportService.is_empty_value(
            remarks
        ):

            document.add_paragraph(
                f"Remarks: {remarks}"
            )

    # =====================================================
    # CT PHASE TABLE
    # =====================================================

    @staticmethod
    def add_phase_test_table(
        document,
        phase_tests
    ):

        valid = [
            item
            for item in (
                phase_tests or []
            )
            if isinstance(
                item,
                dict
            )
        ]

        if not valid:
            return

        preferred = [
            "phase",
            "injected_primary",
            "recorded_secondary",
            "measured_ratio",
            "ratio_error",
            "polarity",
            "polarity_result",
            "result",
        ]

        labels = {
            "phase":
                "Phase",

            "injected_primary":
                "Injected Primary",

            "recorded_secondary":
                "Recorded Secondary",

            "measured_ratio":
                "Measured Ratio",

            "ratio_error":
                "Ratio Error %",

            "polarity":
                "Polarity",

            "polarity_result":
                "Polarity Result",

            "result":
                "Result",
        }

        columns = (
            PanelReportService.get_nonempty_columns(
                valid,
                preferred
            )
        )

        if not columns:
            return

        table = document.add_table(
            rows=1,
            cols=len(columns)
        )

        table.style = "Table Grid"

        for index, key in enumerate(
            columns
        ):

            table.cell(
                0,
                index
            ).text = labels.get(
                key,
                PanelReportService.pretty_label(
                    key
                )
            )

        for phase in valid:

            if all(
                PanelReportService.is_empty_value(
                    phase.get(key)
                )
                for key in columns
            ):
                continue

            cells = table.add_row().cells

            for index, key in enumerate(
                columns
            ):

                value = phase.get(
                    key
                )

                cells[index].text = (
                    ""
                    if PanelReportService.is_empty_value(
                        value
                    )
                    else str(value)
                )

    # =====================================================
    # METER FUNCTION TABLE
    # =====================================================

    @staticmethod
    def add_meter_function_table(
        document,
        functions
    ):

        valid = [
            item
            for item in (
                functions or []
            )
            if isinstance(
                item,
                dict
            )
        ]

        if not valid:
            return

        preferred = [
            "measurement",
            "applied_value",
            "meter_reading",
            "tolerance_percent",
            "error_percent",
            "result",
        ]

        labels = {
            "measurement":
                "Function",

            "applied_value":
                "Applied Value",

            "meter_reading":
                "Recorded Value",

            "tolerance_percent":
                "Tolerance %",

            "error_percent":
                "Error %",

            "result":
                "Result",
        }

        columns = (
            PanelReportService.get_nonempty_columns(
                valid,
                preferred
            )
        )

        if not columns:
            return

        table = document.add_table(
            rows=1,
            cols=len(columns)
        )

        table.style = "Table Grid"

        for index, key in enumerate(
            columns
        ):

            table.cell(
                0,
                index
            ).text = labels.get(
                key,
                PanelReportService.pretty_label(
                    key
                )
            )

        for item in valid:

            if all(
                PanelReportService.is_empty_value(
                    item.get(key)
                )
                for key in columns
            ):
                continue

            cells = table.add_row().cells

            for index, key in enumerate(
                columns
            ):

                value = item.get(
                    key
                )

                cells[index].text = (
                    ""
                    if PanelReportService.is_empty_value(
                        value
                    )
                    else str(value)
                )

    # =====================================================
    # KEY / VALUE TABLE
    # =====================================================

    @staticmethod
    def add_key_value_table(
        document,
        rows
    ):

        rows = (
            PanelReportService.remove_empty_rows(
                rows
            )
        )

        if not rows:
            return

        table = document.add_table(
            rows=1,
            cols=2
        )

        table.style = "Table Grid"

        table.cell(
            0,
            0
        ).text = "Parameter"

        table.cell(
            0,
            1
        ).text = "Value"

        for label, value in rows:

            cells = table.add_row().cells

            cells[0].text = str(
                label
            )

            cells[1].text = str(
                value
            )

    # =====================================================
    # GENERAL HELPERS
    # =====================================================

    @staticmethod
    def get_value(
        obj,
        key,
        default=None
    ):

        if obj is None:
            return default

        if isinstance(
            obj,
            dict
        ):
            return obj.get(
                key,
                default
            )

        return getattr(
            obj,
            key,
            default
        )

    @staticmethod
    def is_empty_value(
        value
    ):

        if value is None:
            return True

        if isinstance(
            value,
            str
        ):

            return (
                value.strip().lower()
                in (
                    "",
                    "null",
                    "none",
                    "n/a",
                    "na",
                    "-"
                )
            )

        return False

    @staticmethod
    def remove_empty_rows(
        rows
    ):

        return [
            (
                label,
                value
            )
            for label, value in (
                rows or []
            )
            if not PanelReportService.is_empty_value(
                value
            )
        ]

    @staticmethod
    def add_summary_row(
        table,
        values
    ):

        if all(
            PanelReportService.is_empty_value(
                value
            )
            for value in values
        ):
            return

        cells = table.add_row().cells

        for index, value in enumerate(
            values
        ):

            cells[index].text = (
                ""
                if PanelReportService.is_empty_value(
                    value
                )
                else str(value)
            )

    @staticmethod
    def get_nonempty_columns(
        rows,
        preferred
    ):

        available = []

        for row in rows:

            for key in row.keys():

                if key not in available:
                    available.append(
                        key
                    )

        columns = [
            key
            for key in preferred
            if key in available
            and any(
                not PanelReportService.is_empty_value(
                    row.get(key)
                )
                for row in rows
            )
        ]

        columns.extend(
            key
            for key in available
            if key not in columns
            and any(
                not PanelReportService.is_empty_value(
                    row.get(key)
                )
                for row in rows
            )
        )

        return columns

    @staticmethod
    def pretty_label(
        value
    ):

        return (
            str(value)
            .replace(
                "_",
                " "
            )
            .replace(
                "-",
                " "
            )
            .strip()
            .title()
        )

    @staticmethod
    def date_only(
        value
    ):

        if value is None:
            return ""

        if isinstance(
            value,
            datetime
        ):
            return value.strftime(
                "%Y-%m-%d"
            )

        text = str(
            value
        ).strip()

        if not text:
            return ""

        if "T" in text:
            return text.split(
                "T",
                1
            )[0]

        if " " in text:
            return text.split(
                " ",
                1
            )[0]

        return text

    @staticmethod
    def normalize_date(
        value
    ):

        if value is None:
            return ""

        if isinstance(
            value,
            datetime
        ):
            return value.strftime(
                "%Y-%m-%d"
            )

        text = str(
            value
        ).strip()

        if not text:
            return ""

        # Common Qt / ISO date representations.
        if "T" in text:
            text = text.split(
                "T",
                1
            )[0]

        elif " " in text:
            text = text.split(
                " ",
                1
            )[0]

        for fmt in (
            "%Y-%m-%d",
            "%d-%m-%Y",
            "%d/%m/%Y",
            "%d.%m.%Y",
            "%Y/%m/%d",
        ):

            try:

                return datetime.strptime(
                    text,
                    fmt
                ).strftime(
                    "%Y-%m-%d"
                )

            except ValueError:
                pass

        return text

    @staticmethod
    def safe_filename_part(
        value
    ):

        if value is None:
            return ""

        text = str(
            value
        ).strip()

        if not text:
            return ""

        invalid = (
            "<",
            ">",
            ":",
            '"',
            "/",
            "\\",
            "|",
            "?",
            "*",
        )

        for character in invalid:
            text = text.replace(
                character,
                "-"
            )

        return text

    @staticmethod
    def find_parent_name(
        node,
        attributes
    ):

        current = node

        for _ in range(4):

            if current is None:
                break

            for attribute in attributes:

                value = getattr(
                    current,
                    attribute,
                    None
                )

                if value:
                    return str(
                        value
                    )

            current = getattr(
                current,
                "parent",
                None
            )

        return ""

    @staticmethod
    def calculate_overall_result(
        protection_tests,
        component_tests
    ):

        results = []

        for test in (
            list(
                protection_tests or []
            )
            +
            list(
                component_tests or []
            )
        ):

            result = str(
                PanelReportService.get_value(
                    test,
                    "result",
                    ""
                )
                or ""
            ).strip().upper()

            if result:
                results.append(
                    result
                )

        if not results:
            return "NOT TESTED"

        if any(
            result == "FAIL"
            for result in results
        ):
            return "FAIL"

        if all(
            result == "PASS"
            for result in results
        ):
            return "PASS"

        return "PARTIALLY TESTED"

    # =====================================================
    # DOCUMENT FORMATTING
    # =====================================================

    @staticmethod
    def configure_document(
        document
    ):

        section = document.sections[0]

        section.top_margin = Inches(
            0.6
        )

        section.bottom_margin = Inches(
            0.6
        )

        section.left_margin = Inches(
            0.6
        )

        section.right_margin = Inches(
            0.6
        )

        normal_style = document.styles[
            "Normal"
        ]

        normal_style.font.name = (
            "Arial"
        )

        normal_style.font.size = Pt(
            9
        )

    @staticmethod
    def add_title(
        document,
        text
    ):

        paragraph = document.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            str(text)
        )

        run.bold = True
        run.font.size = Pt(
            18
        )

    @staticmethod
    def add_heading(
        document,
        text
    ):

        paragraph = document.add_paragraph()

        run = paragraph.add_run(
            str(text)
        )

        run.bold = True
        run.font.size = Pt(
            12
        )

    @staticmethod
    def add_result(
        document,
        result
    ):

        paragraph = document.add_paragraph()

        run = paragraph.add_run(
            f"Result: {result}"
        )

        run.bold = True
        run.font.size = Pt(
            11
        )
