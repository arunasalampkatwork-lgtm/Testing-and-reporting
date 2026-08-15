from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PySide6.QtWidgets import (
    QFileDialog,
    QMessageBox,
)


class AuxRelayReportService:

    def __init__(
        self,
        project_folder
    ):

        self.project_folder = Path(
            project_folder
        )

    # =====================================================
    # GENERATE
    # =====================================================

    def generate_report(
        self,
        record,
        parent=None
    ):

        test_id = record.get(
            "test_id",
            "AUX_RELAY_TEST"
        )

        default_path = (
            self.project_folder
            /
            "reports"
            /
            f"Aux_Relay_Test_{test_id}.docx"
        )

        output_path, _ = (
            QFileDialog.getSaveFileName(

                parent,

                "Save Auxiliary Relay Report",

                str(
                    default_path
                ),

                "Word Document (*.docx)"
            )
        )

        if not output_path:

            return

        output_path = Path(
            output_path
        )

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True
        )

        measurements = (
            record.get(
                "measurements",
                {}
            )
            or
            {}
        )

        document = Document()

        self.configure_document(
            document
        )

        self.add_title(
            document,
            "AUXILIARY RELAY TEST REPORT"
        )

        # =================================================
        # IDENTIFICATION
        # =================================================

        self.add_heading(
            document,
            "RELAY IDENTIFICATION"
        )

        self.add_table(
            document,
            [

                (
                    "Test ID",
                    record.get(
                        "test_id"
                    )
                ),

                (
                    "Test Date",
                    record.get(
                        "test_date"
                    )
                ),

                (
                    "Component ID",
                    record.get(
                        "component_id"
                    )
                ),

                (
                    "Manufacturer",
                    measurements.get(
                        "manufacturer",
                        ""
                    )
                ),

                (
                    "Model",
                    measurements.get(
                        "model",
                        ""
                    )
                ),

                (
                    "Serial Number",
                    measurements.get(
                        "serial_number",
                        ""
                    )
                ),

                (
                    "Coil Voltage",
                    measurements.get(
                        "coil_voltage",
                        ""
                    )
                ),

                (
                    "Contact Configuration",
                    measurements.get(
                        "contact_configuration",
                        ""
                    )
                ),
            ]
        )

        # =================================================
        # COIL
        # =================================================

        self.add_heading(
            document,
            "COIL PICKUP / DROPOUT TEST"
        )

        self.add_table(
            document,
            [

                (
                    "Rated Coil Voltage",
                    measurements.get(
                        "rated_voltage",
                        ""
                    )
                ),

                (
                    "Pickup Voltage",
                    measurements.get(
                        "pickup_voltage",
                        ""
                    )
                ),

                (
                    "Pickup Voltage (%)",
                    measurements.get(
                        "pickup_voltage_percent",
                        ""
                    )
                ),

                (
                    "Dropout Voltage",
                    measurements.get(
                        "dropout_voltage",
                        ""
                    )
                ),

                (
                    "Dropout Voltage (%)",
                    measurements.get(
                        "dropout_voltage_percent",
                        ""
                    )
                ),
            ]
        )

        # =================================================
        # TIMING
        # =================================================

        self.add_heading(
            document,
            "OPERATING TIME"
        )

        self.add_table(
            document,
            [

                (
                    "Expected Pickup Time",
                    measurements.get(
                        "expected_pickup_time",
                        ""
                    )
                ),

                (
                    "Measured Pickup Time",
                    measurements.get(
                        "pickup_time",
                        ""
                    )
                ),

                (
                    "Pickup Time Error (%)",
                    measurements.get(
                        "pickup_time_error",
                        ""
                    )
                ),

                (
                    "Expected Dropout Time",
                    measurements.get(
                        "expected_dropout_time",
                        ""
                    )
                ),

                (
                    "Measured Dropout Time",
                    measurements.get(
                        "dropout_time",
                        ""
                    )
                ),

                (
                    "Dropout Time Error (%)",
                    measurements.get(
                        "dropout_time_error",
                        ""
                    )
                ),
            ]
        )

        # =================================================
        # CONTACT
        # =================================================

        self.add_heading(
            document,
            "CONTACT OPERATION"
        )

        self.add_table(
            document,
            [

                (
                    "Expected Operation",
                    measurements.get(
                        "expected_operation",
                        ""
                    )
                ),

                (
                    "Observed Operation",
                    measurements.get(
                        "observed_operation",
                        ""
                    )
                ),

                (
                    "Functional Result",
                    measurements.get(
                        "functional_result",
                        ""
                    )
                ),
            ]
        )

        # =================================================
        # RESULT
        # =================================================

        self.add_heading(
            document,
            "TEST RESULT"
        )

        paragraph = document.add_paragraph()

        run = paragraph.add_run(
            str(
                record.get(
                    "result",
                    ""
                )
            )
        )

        run.bold = True

        run.font.size = Pt(
            13
        )

        # =================================================
        # REMARKS
        # =================================================

        self.add_heading(
            document,
            "REMARKS"
        )

        document.add_paragraph(
            str(
                record.get(
                    "remarks",
                    ""
                )
                or
                ""
            )
        )

        document.save(
            str(
                output_path
            )
        )

        QMessageBox.information(
            parent,
            "Report Generated",
            (
                "Auxiliary relay report generated "
                "successfully.\n\n"
                f"{output_path}"
            )
        )

        return output_path

    # =====================================================
    # HELPERS
    # =====================================================

    @staticmethod
    def configure_document(
        document
    ):

        section = (
            document.sections[0]
        )

        section.top_margin = Inches(
            0.6
        )

        section.bottom_margin = Inches(
            0.6
        )

        section.left_margin = Inches(
            0.7
        )

        section.right_margin = Inches(
            0.7
        )

        document.styles[
            "Normal"
        ].font.name = "Arial"

        document.styles[
            "Normal"
        ].font.size = Pt(
            9
        )

    @staticmethod
    def add_title(
        document,
        text
    ):

        paragraph = (
            document.add_paragraph()
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            text
        )

        run.bold = True

        run.font.size = Pt(
            17
        )

    @staticmethod
    def add_heading(
        document,
        text
    ):

        paragraph = (
            document.add_paragraph()
        )

        run = paragraph.add_run(
            text
        )

        run.bold = True

        run.font.size = Pt(
            12
        )

    @staticmethod
    def add_table(
        document,
        rows
    ):

        table = document.add_table(
            rows=1,
            cols=2
        )

        table.style = (
            "Table Grid"
        )

        table.rows[0].cells[0].text = (
            "Parameter"
        )

        table.rows[0].cells[1].text = (
            "Value"
        )

        for cell in table.rows[0].cells:

            for run in (
                cell
                .paragraphs[0]
                .runs
            ):

                run.bold = True

        for label, value in rows:

            cells = (
                table.add_row()
                .cells
            )

            cells[0].text = str(
                label
            )

            cells[1].text = (
                ""
                if value is None
                else str(value)
            )