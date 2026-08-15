from pathlib import Path
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches


class ProtectionReportService:
    """Generate a DOCX report from one persisted protection test."""

    def __init__(self, test_service, project_folder):
        self.test_service = test_service
        self.project_folder = Path(project_folder)

    def generate_protection_test_report(self, test_id, output_path=None):
        record = self.test_service.get_test(test_id)
        if record is None:
            raise ValueError(f"Protection test '{test_id}' was not found.")

        if output_path is None:
            output_path = (
                self.project_folder / "reports" /
                f"Protection_Test_{test_id}.docx"
            )

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        context = self._build_context(record)
        document = Document()
        self._configure_document(document)

        self._add_title(document, "PROTECTION RELAY TEST REPORT")
        self._add_test_information(document, context)
        self._add_relay_details(document, context)
        self._add_ct_details(document, context)
        self._add_configuration(document, context)
        self._add_results(document, context)
        self._add_remarks(document, context)
        self._add_signoff(document)

        document.save(str(output_path))
        return output_path

    def _build_context(self, record):
        assets = self._read_json(self.project_folder / "assets.json", [])
        components = self._read_json(self.project_folder / "components.json", [])

        node_map = {
            str(x.get("node_id")): x for x in assets if isinstance(x, dict)
        }
        component_map = {
            str(x.get("component_id")): x
            for x in components if isinstance(x, dict)
        }

        panel = node_map.get(str(record.get("panel_id")), {})
        relay = component_map.get(str(record.get("relay_id")), {})

        hierarchy = []
        current = panel
        seen = set()

        while isinstance(current, dict):
            node_id = current.get("node_id")
            if node_id in seen:
                break
            seen.add(node_id)
            hierarchy.append(current)
            parent_id = current.get("parent_id")
            if parent_id is None:
                break
            current = node_map.get(str(parent_id), {})

        hierarchy.reverse()

        return {
            "record": record,
            "panel": panel,
            "relay": relay,
            "hierarchy": hierarchy,
            "settings": record.get("settings") or {},
            "measurements": record.get("measurements") or {},
            "project_name": self.project_folder.name,
        }

    @staticmethod
    def _read_json(path, default):
        if not path.exists():
            return default
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (OSError, json.JSONDecodeError, TypeError):
            return default

    @staticmethod
    def _configure_document(document):
        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        document.styles["Normal"].font.name = "Arial"
        document.styles["Normal"].font.size = Pt(9)

    @staticmethod
    def _add_title(document, title):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(18)

    def _add_test_information(self, document, c):
        substation = ""
        switchboard = ""
        for node in c["hierarchy"]:
            t = str(node.get("node_type", "")).upper()
            if t == "SUBSTATION":
                substation = node.get("name", "")
            elif t == "SWITCHBOARD":
                switchboard = node.get("name", "")

        rows = [
            ("Project", c["project_name"]),
            ("Substation", substation),
            ("Switchboard", switchboard),
            ("Panel", c["panel"].get("name", "")),
            ("Equipment", c["panel"].get("equipment_name", "")),
            ("Equipment Type", c["panel"].get("equipment_type", "")),
            ("Test ID", c["record"].get("test_id", "")),
            ("Tested Date", c["record"].get("test_date", "")),
        ]
        self._two_column_table(document, "TEST INFORMATION", rows)

    def _add_relay_details(self, document, c):
        relay = c["relay"]
        rows = [
            ("Relay", relay.get("name", "")),
            ("Make", relay.get("manufacturer", "")),
            ("Model No", relay.get("model", "")),
            ("Sr. No", relay.get("serial_number", "")),
            ("Aux. Supply", relay.get("aux_supply", "")),
            ("VT Ratio", relay.get("vt_ratio", "")),
            ("Firmware", relay.get("firmware", "")),
            ("Protection Function", c["record"].get("protection_code", "")),
        ]
        self._two_column_table(document, "RELAY DETAILS", rows)

    def _add_ct_details(self, document, c):
        s = c["settings"]
        rows = [
            ("CT", s.get("ct_name", "")),
            ("CT Ratio", s.get("ct_ratio", "")),
            ("CT Primary", s.get("ct_primary_a", "")),
            ("CT Secondary / In", s.get(
                "ct_secondary_a", s.get("nominal_current_a", "")
            )),
            ("CT Class", s.get("ct_class", "")),
            ("Core", s.get("ct_core", s.get("core", ""))),
        ]
        self._two_column_table(document, "CT DETAILS", rows)

    def _add_configuration(self, document, c):
        excluded = {
            "ct_id", "ct_name", "ct_primary_a", "ct_secondary_a",
            "ct_ratio", "ct_class", "ct_core", "core",
            "nominal_current_a", "nominal_current_unit",
        }
        rows = [
            (self._pretty(k), self._display(v))
            for k, v in c["settings"].items()
            if k not in excluded
        ]
        if rows:
            self._two_column_table(document, "TEST CONFIGURATION", rows)

    def _add_results(self, document, c):
        document.add_heading("TEST RESULTS", level=1)
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, text in enumerate(("Parameter", "Value", "Unit")):
            table.rows[0].cells[i].text = text
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True

        for key, value in c["measurements"].items():
            cells = table.add_row().cells
            cells[0].text = self._pretty(key)
            cells[1].text = self._display(value)
            cells[2].text = self._unit(key)

        p = document.add_paragraph()
        r = p.add_run(f"RESULT: {c['record'].get('result', '')}")
        r.bold = True
        r.font.size = Pt(12)

    def _add_remarks(self, document, c):
        document.add_heading("REMARKS", level=1)
        document.add_paragraph(c["record"].get("remarks", "") or "")

    @staticmethod
    def _add_signoff(document):
        document.add_heading("TESTING / APPROVAL", level=1)
        table = document.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        for i, label in enumerate(("Tested By", "Reviewed By", "Date / Signature")):
            table.rows[i].cells[0].text = label

    def _two_column_table(self, document, title, rows):
        document.add_heading(title, level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "Parameter"
        table.rows[0].cells[1].text = "Value"
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(label)
            cells[1].text = self._display(value)

    @staticmethod
    def _display(value):
        if value is None:
            return ""
        if isinstance(value, (dict, list)):
            return json.dumps(value, indent=2)
        return str(value)

    @staticmethod
    def _pretty(value):
        return str(value).replace("_", " ").strip().title()

    @staticmethod
    def _unit(key):
        k = str(key).lower()
        if "xin" in k or "_x_in" in k:
            return "xIn"
        if k.endswith("_a") or "_current_a" in k:
            return "A"
        if "time" in k or "duration" in k:
            return "s"
        if "angle" in k:
            return "deg"
        if "frequency" in k:
            return "Hz"
        if "rocof" in k:
            return "Hz/s"
        if "percent" in k or "error" in k:
            return "%"
        return ""
